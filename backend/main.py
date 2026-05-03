from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
# REMOVE THE DOTS HERE:
from parser import extract_text
from ai_engine import analyze_cv_with_ai
from fastapi.responses import FileResponse
from docx import Document
import os
app = FastAPI()

# Allow your future React frontend to connect
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.post("/analyze")
async def analyze(file: UploadFile = File(...), job_description: str = Form(...)):
    content = await file.read()
    text = extract_text(content, file.filename)
    
    if not text:
        return {"error": "Format not supported"}
    
    result = analyze_cv_with_ai(text, job_description)
    return result

@app.post("/download-cv")
async def download_cv(data: dict):
    # data will contain the 'improved_bullets' from the AI
    bullets = data.get("improved_bullets", [])
    
    # 1. Create a Word Document
    doc = Document()
    doc.add_heading('Rewritten Professional Experience', 0)
    
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')
    
    # 2. Save locally temporarily
    file_path = "rewritten_cv.docx"
    doc.save(file_path)
    
    # 3. Return the file to the user
    return FileResponse(
        path=file_path, 
        filename="Rewritten_CV.docx", 
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)